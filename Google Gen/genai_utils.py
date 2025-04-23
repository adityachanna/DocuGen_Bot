import os
import pathlib
import re 
from google import genai
from google.genai.types import Tool, GenerateContentConfig, GoogleSearch
from google.genai import types
def load_genai_client():
    api_key = os.getenv('GEMINI_API_KEY')
    if not api_key:
        raise ValueError("GEMINI_API_KEY environment variable not set.")
    os.environ['GOOGLE_API_KEY'] = api_key
    return genai.Client()

def process_pdf(file_path):
    """Reads a PDF file and returns its content as bytes."""
    try:
        return pathlib.Path(file_path).read_bytes()
    except FileNotFoundError:
        print(f"Error: File not found at {file_path}")
        raise
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        raise
def generate_document(client, file_bytes, prompt, model_id="gemini-2.0-flash", use_google_search=True):
    """Generates text content from PDF bytes using Google GenAI, optionally with Google Search grounding."""
    gen_config = None 
    tools = None
    if use_google_search:
        try:
            search_tool = {
                'google_search': {}
            }
            gen_config = GenerateContentConfig(
                tools=[search_tool],
                response_modalities=["TEXT"]
            )
        except Exception as e:
            print(f"Warning: Could not configure Google Search: {e}")
            print("Continuing without search capability...")
            use_google_search = False
    try:
        contents_list = [
            types.Part.from_bytes(
                data=file_bytes,
                mime_type='application/pdf',
            ),
            prompt
        ]
        response = client.models.generate_content(
            model=model_id,
            contents=contents_list,
            config=gen_config 
        )
        if hasattr(response, 'candidates') and response.candidates:
            print("Grounding data:", response.candidates[0].grounding_metadata or "No grounding data")
        return response.text 
    except Exception as e:
        print(f"Error during GenAI content generation: {e}")
        raise
def _add_formatted_runs_docx(paragraph, text_line):
    """Adds runs to a DOCX paragraph, applying bold/italic based on markdown."""
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', text_line)
    if heading_match:
        paragraph.clear()
        heading_level = len(heading_match.group(1))
        heading_text = heading_match.group(2)
        heading = paragraph.part.document.add_heading(level=heading_level)
        heading.text = heading_text
        return
    current_pos = 0
    for match in re.finditer(r'(\*\*.*?\*\*)|(\*.*?\*)', text_line):
        start, end = match.span()
        if start > current_pos:
            paragraph.add_run(text_line[current_pos:start])
        matched_text = match.group(0)
        if matched_text.startswith('**') and matched_text.endswith('**'):
            run = paragraph.add_run(matched_text[2:-2])
            run.bold = True
        elif matched_text.startswith('*') and matched_text.endswith('*'):
            run = paragraph.add_run(matched_text[1:-1])
            run.italic = True
        current_pos = end
    if current_pos < len(text_line):
        paragraph.add_run(text_line[current_pos:])
def _write_formatted_line_pdf(pdf, text_line, line_height=5, font_name="Arial", font_size=12):
    """Writes a line to the PDF, applying bold/italic styles and handling basic wrapping."""
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', text_line)
    if heading_match:
        heading_level = len(heading_match.group(1))
        heading_text = heading_match.group(2)
        heading_size = 20 - (heading_level * 2)
        pdf.set_font(font_name, style='B', size=heading_size)
        pdf.write(line_height, heading_text.encode('latin-1', 'replace').decode('latin-1'))
        pdf.ln(line_height * 2)  
        return
    processed_line = text_line
    processed_line = processed_line.replace('\\*', '§ESCAPED-ASTERISK§')
    processed_line = processed_line.replace('/*', '§SLASH-ASTERISK§')
    processed_line = re.sub(r'(\*\*[^*]+\*\*)|(\*[^*]+\*)', 
                          lambda m: m.group(0).replace('*', '§FORMAT-ASTERISK§'), 
                          processed_line)
    
    processed_line = processed_line.replace('*', '\\*')
    processed_line = processed_line.replace('§FORMAT-ASTERISK§', '*')
    processed_line = processed_line.replace('§ESCAPED-ASTERISK§', '\\*')
    processed_line = processed_line.replace('§SLASH-ASTERISK§', '/*')
  
    current_pos = 0
    pdf.set_font(font_name, style='', size=font_size)
    for match in re.finditer(r'(\*\*[^*]+\*\*)|(\*[^*]+\*)', processed_line):
        start, end = match.span()
        if start > current_pos:
            pdf.set_font(font_name, style='', size=font_size)
            prefix_text = processed_line[current_pos:start].encode('latin-1', 'replace').decode('latin-1')
            pdf.write(line_height, prefix_text)

        matched_text = match.group(0)
        style_to_set = ''
        inner_text = ''
        
        if matched_text.startswith('**') and matched_text.endswith('**'):
            style_to_set = 'B' 
            inner_text = matched_text[2:-2]
        elif matched_text.startswith('*') and matched_text.endswith('*'):
            style_to_set = 'I'
            inner_text = matched_text[1:-1]

        pdf.set_font(font_name, style=style_to_set, size=font_size)
        formatted_text = inner_text.encode('latin-1', 'replace').decode('latin-1')
        pdf.write(line_height, formatted_text)

        current_pos = end
    if current_pos < len(processed_line):
        pdf.set_font(font_name, style='', size=font_size)
        suffix_text = processed_line[current_pos:].encode('latin-1', 'replace').decode('latin-1')
        pdf.write(line_height, suffix_text)

    pdf.ln(line_height)
def save_output_document(content, output_format, output_path):
    """Saves the generated content, applying basic markdown formatting (bold/italic)."""
    output_path = pathlib.Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    lines = content.splitlines() 

    if output_format.lower() == 'pdf':
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=14) 

            for line in lines:
                if line.strip(): 
                     _write_formatted_line_pdf(pdf, line)
                else: 
                    pdf.ln(5) 

            pdf.output(str(output_path))
            print(f"Generated PDF saved to: {output_path}")
        except ImportError:
            print("Error: fpdf2 library not found. Cannot save as PDF.")
            print("Install it using: pip install fpdf2")
            # Fallback: Save as plain text
            txt_path = output_path.with_suffix('.txt')
            print(f"Saving content as plain text instead: {txt_path}")
            txt_path.write_text(content, encoding='utf-8')
        except Exception as e:
            print(f"Error saving PDF file: {e}")

    elif output_format.lower() == 'docx':
        try:
            from docx import Document
            doc = Document()

            for line in lines:
               
                 p = doc.add_paragraph()
                 if line.strip():
                     _add_formatted_runs_docx(p, line)
                
            doc.save(output_path)
            print(f"Generated DOCX saved to: {output_path}")
        except ImportError:
             print("Error: python-docx library not found. Cannot save as DOCX.")
             print("Install it using: pip install python-docx")
        except Exception as e:
            print(f"Error saving DOCX file: {e}")
    else:
        print(f"Unsupported output format '{output_format}'. Saving as .txt")
        txt_path = output_path.with_suffix('.txt')
        txt_path.write_text(content, encoding='utf-8') 
